# -*- coding: utf-8 -*-
"""
HIGGS projesi icin ozgun bir Word (.docx) raporu uretir (Turkce karakterli).

Gorseller icin sayfada cerceveli "yer tutucu" kutular birakir; her kutunun
altinda hangi grafik dosyasinin eklenecegi ve sekil aciklamasi belirtilir.
Calistirma:  python generate_report.py
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "MehmetSen_MakineOgrenmesi_Final_Rapor.docx")

FIG_DIR = "outputs/figures"


# ---------------------------------------------------------------------------
# Yardimci fonksiyonlar
# ---------------------------------------------------------------------------

def set_cell_background(cell, color_hex):
    """Bir tablo hucresine arka plan rengi verir."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_image_placeholder(doc, fig_number, caption, filename, n_lines=6):
    """Cerceveli bir gorsel yer tutucu kutusu + sekil aciklamasi ekler."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    cell = table.cell(0, 0)
    set_cell_background(cell, 'F2F2F2')

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[ GÖRSEL ALANI ]")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Buraya ekleyin:  {FIG_DIR}/{filename}")
    r2.italic = True
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    for _ in range(n_lines):
        cell.add_paragraph()

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(f"Şekil {fig_number}: {caption}")
    cr.italic = True
    cr.font.size = Pt(9)
    doc.add_paragraph()


def add_table_from_data(doc, headers, rows, caption=None, table_number=None):
    """Basliklar + satirlardan bicimli bir tablo olusturur."""
    if caption and table_number:
        cap = doc.add_paragraph()
        cr = cap.add_run(f"Tablo {table_number}: {caption}")
        cr.bold = True
        cr.font.size = Pt(10)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


# ---------------------------------------------------------------------------
# Rapor icerigi
# ---------------------------------------------------------------------------

def build_report():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ----- KAPAK -----
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("HIGGS Veri Seti Üzerinde Makine Öğrenmesi Hattı")
    tr.bold = True
    tr.font.size = Pt(22)

    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t2r = t2.add_run("Özellik Seçimi ve Hiperparametre Optimizasyonu")
    t2r.font.size = Pt(16)
    t2r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    for _ in range(6):
        doc.add_paragraph()

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ir = info.add_run("Makine Öğrenmesi - Final Ödevi\n\n"
                      "Mehmet Şen\n"
                      "Öğrenci No: 254312024\n\n"
                      "Üsküdar Üniversitesi")
    ir.font.size = Pt(13)

    doc.add_page_break()

    # ----- 1. GIRIS -----
    add_heading(doc, "1. Giriş ve Amaç", 1)
    doc.add_paragraph(
        "Bu çalışmada, yüksek enerji fiziği alanında Higgs bozonu üretim "
        "süreçlerini sinyal ve arka plan olayları olarak ayırt etmeyi amaçlayan "
        "HIGGS veri seti üzerinde uçtan uca bir makine öğrenmesi hattı (pipeline) "
        "tasarlanmıştır. Çalışmanın iki temel bileşeni; filtre tabanlı özellik "
        "seçimi (feature selection) ve nested cross-validation ile hiperparametre "
        "optimizasyonudur (hyperparameter tuning)."
    )
    doc.add_paragraph(
        "Hedef; K-En Yakın Komşu (KNN), Destek Vektör Makinesi (SVM), Çok Katmanlı "
        "Algılayıcı (MLP) ve XGBoost algoritmalarını aynı koşullar altında eğiterek "
        "performanslarını Accuracy, Precision, Recall, F1 ve ROC-AUC metrikleri ile "
        "karşılaştırmak ve en başarılı model-temsil kombinasyonunu belirlemektir."
    )

    # ----- 2. VERI SETI -----
    add_heading(doc, "2. Veri Seti", 1)
    doc.add_paragraph(
        "HIGGS veri seti yaklaşık 11 milyon örnek ve 28 sayısal özellik içermektedir. "
        "Özelliklerin ilk 21'i düşük seviyeli kinematik büyüklüklerden (parçacıkların "
        "momentum, açı ve etiket bilgileri), son 7'si ise bu büyüklüklerden türeyen "
        "yüksek seviyeli kütle değişkenlerinden (m_jj, m_jjj, m_lv, m_jlv, m_bb, "
        "m_wbb, m_wwbb) oluşur. Etiket ikili (binary) olup 1 = sinyal, 0 = arka plandır."
    )
    doc.add_paragraph(
        "Hesaplama maliyetini yönetebilmek için tüm veri setinden tekrar üretilebilir "
        "bir tohum (random_state=42) ile rastgele 100.000 örnek seçilmiştir. Veri, "
        "büyük gzip dosyasından parça parça (chunk) okunarak örneklendirilmiş ve "
        "sonraki çalışmalar için parquet formatında önbelleğe alınmıştır."
    )

    # ----- 3. YONTEM -----
    add_heading(doc, "3. Yöntem", 1)

    add_heading(doc, "3.1 Veri Ön İşleme", 2)
    doc.add_paragraph(
        "Aykırı değer analizi IQR (Interquartile Range) yöntemi ile yapılmıştır. "
        "Her özellik için Q1 ve Q3 çeyrekleri hesaplanmış, IQR = Q3 - Q1 olmak üzere "
        "[Q1 - 1.5*IQR, Q3 + 1.5*IQR] aralığı dışındaki değerler aykırı kabul "
        "edilmiştir. Örnek sayısını ve sınıf dengesini korumak amacıyla aykırı "
        "değerler silinmek yerine ilgili sınır değerlere kırpılmıştır (winsorization). "
        "Ardından tüm özellikler MinMaxScaler ile [0, 1] aralığına ölçeklenmiştir. "
        "Veri sızıntısını (data leakage) önlemek için ölçekleme her çapraz doğrulama "
        "katmanında yalnızca eğitim verisine uyarlanmıştır."
    )

    add_heading(doc, "3.2 Özellik Seçimi", 2)
    doc.add_paragraph(
        "Filtre tabanlı özellik seçimi için ANOVA F-skoru (f_classif) kullanılmış ve "
        "en yüksek skora sahip 15 özellik seçilmiştir. Bu yöntem, her özellik ile "
        "hedef değişken arasındaki istatistiksel ilişkiyi ölçerek modelden bağımsız, "
        "hızlı bir sıralama sağlar."
    )

    add_heading(doc, "3.3 Nested Cross-Validation (Flowchart A ve B)", 2)
    doc.add_paragraph(
        "Model seçimi ve değerlendirme yanlılığını önlemek için iç içe (nested) "
        "çapraz doğrulama uygulanmıştır. Dış döngü (outer loop) 5 katlı, iç döngü "
        "(inner loop) 3 katlıdır. Dış döngü bağımsız test performansını ölçerken, iç "
        "döngü en iyi konfigürasyonu seçer:"
    )
    doc.add_paragraph(
        "Flowchart A: İç döngüde farklı öznitelik seçim kombinasyonları (seçilen "
        "özellik sayısı k = 10, 15, 20, 28) denenerek en iyi öznitelik alt kümesi "
        "belirlenir.", style='List Bullet')
    doc.add_paragraph(
        "Flowchart B: İç döngüde farklı hiperparametre kombinasyonları denenerek her "
        "model için en iyi hiperparametreler belirlenir.", style='List Bullet')

    add_image_placeholder(doc, "1", "Flowchart A - İç döngüde öznitelik seçimi akış şeması "
                          "(ödev dokümanındaki Figure 1).", "FlowchartA.png")
    add_image_placeholder(doc, "2", "Flowchart B - İç döngüde hiperparametre optimizasyonu akış şeması "
                          "(ödev dokümanındaki Figure 2).", "FlowchartB.png")

    add_heading(doc, "3.4 Modeller ve Hiperparametre Aralıkları", 2)
    add_table_from_data(
        doc,
        ["Model", "Inner CV'de Denenen Hiperparametreler"],
        [
            ["KNN", "n_neighbors = {3, 5, 7, 9, 11}"],
            ["SVM", "C = {0.1, 1, 10}, kernel = {linear, rbf}"],
            ["MLP", "hidden_layer_sizes = {(50,), (100,)}, activation = {relu, tanh}"],
            ["XGBoost", "n_estimators = {100, 200}, max_depth = {3, 6}, learning_rate = {0.1, 0.3}"],
        ],
        caption="Modeller ve inner CV hiperparametre aralıkları.", table_number="1",
    )

    add_heading(doc, "3.5 Değerlendirme Metrikleri", 2)
    doc.add_paragraph(
        "Modeller Accuracy, Precision, Recall, F1 ve ROC-AUC metrikleri ile "
        "değerlendirilmiştir. Ayrıca OVA (One-vs-All) yaklaşımı ile her sınıf (sinyal "
        "ve arka plan) için ROC eğrileri çizilmiş ve AUC skorları yorumlanmıştır."
    )

    # ----- 4. BULGULAR -----
    add_heading(doc, "4. Bulgular", 1)

    add_heading(doc, "4.1 Aykırı Değer Analizi", 2)
    doc.add_paragraph(
        "IQR analizine göre en yüksek aykırı değer oranları, yüksek seviyeli kütle "
        "değişkenlerinde gözlemlenmiştir. Aşağıdaki tabloda en yüksek aykırı değer "
        "oranına sahip özellikler özetlenmiştir."
    )
    add_table_from_data(
        doc,
        ["Özellik", "Aykırı Değer Sayısı", "Oran (%)"],
        [
            ["m_lv", "20.188", "20.19"],
            ["m_jj", "13.780", "13.78"],
            ["m_jjj", "7.536", "7.54"],
            ["m_wbb", "6.351", "6.35"],
            ["m_bb", "6.130", "6.13"],
            ["m_wwbb", "5.867", "5.87"],
            ["m_jlv", "5.075", "5.08"],
        ],
        caption="IQR yöntemine göre en çok aykırı değere sahip 7 özellik.", table_number="2",
    )
    add_image_placeholder(doc, "3", "Ölçeklemeden önce özelliklerin box-plot dağılımları.",
                          "boxplots_before_scaling.png")
    add_image_placeholder(doc, "4", "IQR kırpma (winsorization) sonrası box-plot dağılımları.",
                          "boxplots_after_clipping.png")

    add_heading(doc, "4.2 Özellik Seçimi Sonuçları", 2)
    doc.add_paragraph(
        "ANOVA F-skoruna göre en ayırt edici özellikler, fizik açısından da anlamlı "
        "olan yüksek seviyeli kütle değişkenleri (m_bb, m_wwbb, m_wbb) ile kayıp "
        "enerji büyüklüğü (missing_energy_magnitude) olmuştur. Açı (phi) ve eta "
        "tabanlı değişkenler ise en düşük skorlara sahiptir."
    )
    add_table_from_data(
        doc,
        ["Sıra", "Özellik", "ANOVA F-skoru"],
        [
            ["1", "m_bb", "2017.58"],
            ["2", "m_wwbb", "1898.06"],
            ["3", "missing_energy_magnitude", "836.05"],
            ["4", "m_wbb", "449.05"],
            ["5", "jet1_pt", "325.57"],
            ["6", "jet2_b_tag", "236.17"],
            ["7", "jet4_pt", "154.82"],
            ["8", "lepton_pT", "100.62"],
            ["9", "m_jjj", "86.37"],
            ["10", "jet2_pt", "48.16"],
            ["11", "jet3_b_tag", "47.64"],
            ["12", "m_jj", "24.63"],
            ["13", "jet4_b_tag", "21.87"],
            ["14", "jet3_pt", "13.78"],
            ["15", "jet1_b_tag", "13.44"],
        ],
        caption="ANOVA F-skoruna göre seçilen en iyi 15 özellik.", table_number="3",
    )
    add_image_placeholder(doc, "5", "Tüm özelliklerin ANOVA F-skorları; seçilen en iyi 15 vurgulanmıştır.",
                          "feature_importance_scores.png")

    add_heading(doc, "4.3 Model Performansları - Flowchart B (Hiperparametre Optimizasyonu)", 2)
    doc.add_paragraph(
        "Aşağıdaki tabloda her modelin 5 katlı dış döngüdeki ortalama metrik değerleri "
        "(ortalama ± standart sapma) verilmiştir."
    )
    add_table_from_data(
        doc,
        ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"],
        [
            ["KNN", "0.601 ± 0.015", "0.608 ± 0.014", "0.697 ± 0.016", "0.649 ± 0.011", "0.638 ± 0.015"],
            ["SVM", "0.669 ± 0.007", "0.666 ± 0.009", "0.754 ± 0.024", "0.707 ± 0.009", "0.728 ± 0.008"],
            ["MLP", "0.680 ± 0.004", "0.690 ± 0.017", "0.721 ± 0.043", "0.704 ± 0.012", "0.742 ± 0.009"],
            ["XGBoost", "0.696 ± 0.013", "0.710 ± 0.014", "0.721 ± 0.017", "0.715 ± 0.012", "0.766 ± 0.012"],
        ],
        caption="Flowchart B - Nested CV ortalama performans metrikleri.", table_number="4",
    )
    add_image_placeholder(doc, "6", "Tüm modellerin ROC eğrileri karşılaştırması (Flowchart B).",
                          "roc_curves_all_models_flow_B.png")
    add_image_placeholder(doc, "7", "Modellerin metrik karşılaştırması (Flowchart B).",
                          "metric_comparison_flow_B.png")

    add_heading(doc, "4.4 Model Performansları - Flowchart A (Öznitelik Seçimi)", 2)
    add_table_from_data(
        doc,
        ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC"],
        [
            ["KNN", "0.610 ± 0.014", "0.618 ± 0.014", "0.693 ± 0.017", "0.653 ± 0.011", "0.645 ± 0.012"],
            ["SVM", "0.668 ± 0.015", "0.664 ± 0.016", "0.759 ± 0.016", "0.708 ± 0.011", "0.732 ± 0.011"],
            ["MLP", "0.683 ± 0.008", "0.692 ± 0.014", "0.727 ± 0.018", "0.709 ± 0.007", "0.747 ± 0.012"],
            ["XGBoost", "0.686 ± 0.012", "0.701 ± 0.009", "0.711 ± 0.019", "0.706 ± 0.013", "0.756 ± 0.008"],
        ],
        caption="Flowchart A - Nested CV ortalama performans metrikleri.", table_number="5",
    )
    add_image_placeholder(doc, "8", "Tüm modellerin ROC eğrileri karşılaştırması (Flowchart A).",
                          "roc_curves_all_models_flow_A.png")

    add_heading(doc, "4.5 Seçilen En İyi Hiperparametreler (Flowchart B)", 2)
    add_table_from_data(
        doc,
        ["Model", "Dış Döngüde Seçilen Tipik Hiperparametreler"],
        [
            ["KNN", "n_neighbors = 9-11"],
            ["SVM", "C = 1, kernel = rbf (tüm katmanlar)"],
            ["MLP", "activation = relu, hidden_layer_sizes = (50,) / (100,)"],
            ["XGBoost", "learning_rate = 0.1, max_depth = 3, n_estimators = 100"],
        ],
        caption="Inner CV ile seçilen en iyi hiperparametreler.", table_number="6",
    )
    doc.add_paragraph(
        "Flowchart A'da iç döngü çoğunlukla daha geniş öznitelik kümesini (k = 28) "
        "seçmiştir; ancak 15 özellik ile elde edilen performansa kıyasla kazanç "
        "marjinaldir. Bu durum, seçilen 15 özelliğin ayırt edici bilginin büyük "
        "kısmını taşıdığını göstermektedir."
    )

    add_heading(doc, "4.6 OVA ROC Eğrileri", 2)
    doc.add_paragraph(
        "İkili sınıflandırma probleminde OVA yaklaşımı ile her sınıf (sinyal ve arka "
        "plan) için ROC eğrileri çizilmiştir. İkili durumda iki sınıfın eğrileri "
        "simetriktir ve aynı AUC değerini verir."
    )
    add_image_placeholder(doc, "9", "XGBoost için OVA ROC eğrileri (en başarılı model).",
                          "roc_ova_XGBoost_flow_B.png")

    # ----- 5. TARTISMA -----
    add_heading(doc, "5. Tartışma ve Yorum", 1)
    doc.add_paragraph(
        "Sonuçlar, ağaç tabanlı XGBoost modelinin tüm metriklerde diğerlerinden üstün "
        "olduğunu (ROC-AUC = 0.766) göstermektedir. Bunu sırasıyla MLP (0.742) ve SVM "
        "(0.728) izlemektedir. KNN ise 0.638 ROC-AUC ile en düşük performansı "
        "sergilemiştir; bu beklenen bir sonuçtur, çünkü mesafe tabanlı yöntemler "
        "yüksek boyutlu ve gürültülü verilerde zayıflar."
    )
    doc.add_paragraph(
        "XGBoost'un üstünlüğü; özellikler arasındaki doğrusal olmayan etkileşimleri "
        "ve fizik değişkenleri arasındaki karmaşık ilişkileri ağaç yapılarıyla "
        "yakalayabilmesinden kaynaklanır. Ayrıca seçilen hiperparametrelerin (sığ "
        "ağaçlar: max_depth=3, learning_rate=0.1) tutarlı olması, modelin aşırı "
        "öğrenmeye (overfitting) karşı dengeli kaldığını göstermektedir."
    )
    doc.add_paragraph(
        "Özellik seçimi analizi, ayırt edici gücün büyük ölçüde yüksek seviyeli kütle "
        "değişkenlerinde (m_bb, m_wwbb, m_wbb, m_jjj) ve kayıp enerjide toplandığını "
        "ortaya koymuştur. Bu, fizik literatürü ile uyumludur: bu türetilmiş "
        "değişkenler, ham kinematik büyüklüklerden daha fazla ayırt edici bilgi taşır."
    )

    # ----- 6. SONUC -----
    add_heading(doc, "6. Sonuç", 1)
    doc.add_paragraph(
        "Bu çalışmada HIGGS veri setinden seçilen 100.000 örnek üzerinde tam bir "
        "makine öğrenmesi hattı uygulanmıştır. IQR tabanlı aykırı değer işleme ve "
        "MinMax ölçekleme ile veri hazırlanmış, ANOVA F-skoru ile en iyi 15 özellik "
        "seçilmiş ve nested cross-validation ile dört model adil bir şekilde "
        "karşılaştırılmıştır. En başarılı model-temsil kombinasyonu, ANOVA ile "
        "seçilen özellikler üzerinde eğitilen XGBoost (ROC-AUC = 0.766) olmuştur. "
        "Bu sonuç, yapılandırılmış tablo verilerinde gradyan artırımlı ağaç "
        "yöntemlerinin gücünü bir kez daha doğrulamaktadır."
    )

    # ----- 7. KAYNAKCA -----
    add_heading(doc, "7. Kaynaklar ve Proje Bağlantısı", 1)
    doc.add_paragraph("HIGGS Dataset - UCI Machine Learning Repository: "
                      "https://archive.ics.uci.edu/ml/datasets/HIGGS", style='List Bullet')
    doc.add_paragraph("Proje GitHub deposu: [GitHub bağlantınızı buraya ekleyin]",
                      style='List Bullet')
    doc.add_paragraph("Kullanılan kütüphaneler: scikit-learn, XGBoost, pandas, NumPy, "
                      "Matplotlib, Seaborn.", style='List Bullet')

    doc.save(OUT_PATH)
    print("Rapor olusturuldu:", OUT_PATH)


if __name__ == "__main__":
    build_report()
